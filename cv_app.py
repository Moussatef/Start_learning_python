from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


docx = Document()


class Person:

    def __init__(self, first_name, last_name, age, gender, phone, address,
                 mail):
        self.first_name = first_name
        self.last_name = last_name
        self.age = age
        self.gender = gender
        self.phone = phone
        self.address = address
        self.mail = mail


speak("What is your First Name")
first_name = input("What is your First Name : ")
speak("Helo " + first_name + " can you give me your last name")
last_name = input("What is your Last Name : ")
speak("and what is your email address")
mail = input("What is your Email : ")
speak("Please give your age ")
age = input("What is your Age : ")

speak("Your complet name is " + first_name + " " + last_name +
      " your email is " + mail + " your age is " + age)

while int(age) < 0 and int(age) > 80:
    print("age invalid")
    age = input("What is your age : ")

gender = input("What is your Gender [Male or female] : ")
while gender.lower() != "male" and gender.lower() != "female":
    print("gender invalid")
    gender = input("Enter your gender : ")

phone = input("What is your Phone Number : ")
address = input("What is your Address : ")

person = Person(first_name, first_name, age, gender, phone, address, mail)

docx.add_picture("pictures/QTR11058.jpg", width=Inches(1.4))
docx.add_heading("Personal Information ")
docx.add_paragraph(person.first_name.upper() + " " +
                   person.last_name.upper()).bold = True
docx.add_paragraph("Adrress : " + person.address)
docx.add_paragraph("Age : " + str(person.age))
docx.add_paragraph("Gender : " + person.gender)

docx.add_heading("Contact Information")
contact = "Email : " + person.mail + "\nPhone : " + person.phone

docx.add_paragraph(contact)

# Add experience
docx.add_heading("EXPERIENCES")

while True:

    company = input("Company name : ")
    from_d = input("Start date work : ")
    end_d = input("End date work : ")
    description = input("Description on your experience : ")

    docx.add_paragraph(company.capitalize()).bold = True
    docx.add_paragraph(str(from_d) + " - " + str(end_d)).italic = True
    docx.add_paragraph(description.capitalize())

    more_exp = input("Are you want to add more experience YES or NO ? : ")

    if more_exp.upper == "YES":
        continue
    else:
        break

docx.save("import_result/cv_1.docx")
